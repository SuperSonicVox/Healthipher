from docx import Document
from docx.shared import Inches

def records_in(records, filepath):
    document = Document()
    document.add_heading('結果報表', 0)
    document.add_heading('', level=1)
    if "dropna" in records:
        document.add_paragraph('清除空值', style='Intense Quote')
    if "replace_na" in records:
        document.add_paragraph('以平均值替代空值', style='Intense Quote')
    if "delete_col" in records:
        i = records.index("delete_col")
        content = '刪除行，行名稱為：' + records[i+1]
        document.add_paragraph(content, style='Intense Quote')
    if "select_var" in records:
        i = records.index("select_var")
        content = '進行變數類別選擇\n類別變數：'
        for ele in records[i+1]:
            content = content + ele
        content = content + '\n連續變數：'
        for ele in records[i+2]:
            content = content + ele
        document.add_paragraph(content, style='Intense Quote')
    if "cut" in records:
        document.add_paragraph('進行資料分組，分組變數為：', style='Intense Quote')
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

#執行演算法運算前，系統已自動將文字內容轉換為數值


    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    filepath = filepath + '/result.docx'
    document.save(filepath)