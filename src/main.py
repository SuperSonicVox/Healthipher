import os, shutil
import sys
import csv
import pandas as pd
import numpy as np
from PyQt5 import QtCore, QtGui, QtWidgets, uic
from PyQt5.QtWidgets import QFileDialog, QTableView, QMessageBox, QLabel, QVBoxLayout, QFormLayout, QGroupBox, QCheckBox, QPushButton
from PyQt5.QtGui import QPixmap
from PyQt5.QtCore import Qt
from docx import Document
from matplotlib import pyplot as plt

import tableView, des_sta, pre_processing, checkList, logistic, svm
from testML import cart, knn, rf

class UI(QtWidgets.QMainWindow):
    def __init__(self):
        super(UI, self).__init__()
        uic.loadUi("MainWindow.ui", self)
        self.setWindowTitle("Healthipher")
        #main
        self.btn_upload.clicked.connect(self.load_data) #當load data的按鈕按下去的時候會發生甚麼事
        self.reset_btn.clicked.connect(self.reset)
        #tab_2
        self.null_img.setScaledContents(True)
        self.dropna_btn.clicked.connect(self.dropna)
        self.delete_btn.clicked.connect(self.delete)
        self.replacena_btn.clicked.connect(self.replace_na)
        #tab_3
        self.var_select_btn.clicked.connect(self.send_selected)
        self.var_select_btn.clicked.connect(self.renewal)
        self.var_select_btn.clicked.connect(self.show_des)
        self.cut_btn.clicked.connect(self.cut)
        self.cut_btn.clicked.connect(self.show_cate_plot)
        self.cancel_btn.clicked.connect(self.cancel_cut)
        self.send_cut_btn.clicked.connect(self.send_cut)
        #tab_5
        self.algorithm_connect_btn.clicked.connect(self.convert)
        self.algorithm_connect_btn.clicked.connect(self.cart_show)
        self.algorithm_connect_btn.clicked.connect(self.knn_show)
        self.algorithm_connect_btn.clicked.connect(self.rf_show)
        self.algorithm_connect_btn.clicked.connect(self.svm_show)
        self.algorithm_connect_btn.clicked.connect(self.logistic_show)
        self.para_clear_btn.clicked.connect(self.parameter_clear)
        #tab_6
        self.cart_con.setScaledContents(True)
        self.cart_roc.setScaledContents(True)
        self.knn_con.setScaledContents(True)
        self.knn_roc.setScaledContents(True)
        self.rf_con.setScaledContents(True)
        self.rf_roc.setScaledContents(True)
        self.svm_con.setScaledContents(True)
        self.svm_roc.setScaledContents(True)
        self.logistic_con.setScaledContents(True)
        self.logistic_roc.setScaledContents(True)
        self.csv_download_btn.clicked.connect(self.download_csv)
        self.download_btn.clicked.connect(self.download_file)


        self.show()

    #tab1
    def load_data(self):
        self.document = Document()
        self.document.add_heading('結果報表', 0)
        self.document.add_heading('一、前處理步驟', level=1)
        self.fileName1, filetype = QFileDialog.getOpenFileName(self, "./", "(*csv);; (*.csv)")#最後兩個參數表示只允許顯示 選擇csv檔
        if self.fileName1 != '':
            self.label_1.setText(self.fileName1)#選好檔案之後把"檔案名稱"換成檔案的名字
            self.label_1.setWordWrap(True)

            self.df = pd.read_csv(self.fileName1)
            self.categorial_lst = []
            self.conti_lst = []

            self.algorithm_empty()


            model = tableView.pandasModel(self.df)
            self.dataTable.setModel(model)

            self.var_selection()
            self.renewal()

    #tab2
    def dropna(self):
        self.df.dropna(inplace=True)
        self.document.add_paragraph('清除空值', style='Intense Quote')
        self.renewal()

    def delete(self):
        delete_col = str(self.col_value_delete.currentText())
        self.df = self.df.drop(delete_col, 1)

        if delete_col in self.conti_lst:
            self.conti_lst.remove(delete_col)
        elif delete_col in self.categorial_lst:
            self.categorial_lst.remove(delete_col)

        self.var_selection()
        content = '刪除行，行名稱為：' + delete_col
        self.document.add_paragraph(content, style='Intense Quote')
        self.renewal()

    def replace_na(self):
        self.df.fillna(self.df.mean(), inplace=True)
        self.document.add_paragraph('以平均值替代空值', style='Intense Quote')
        self.renewal()

    #tab3
    def checklist_show(self):
        checkList.nomi_or_conti(self.df)

    def var_selection(self):
        self.formLayout = QFormLayout()
        groupBox = QGroupBox()

        i = self.df.shape[1]
        header = []
        contain = []
        self.df.columns = self.df.columns.str.strip()
        contains = self.df.loc[0]
        for string in self.df.columns:
            header.append(string)
        for string in contains:
            contain.append(string)

        for j in range(i):
            checkbox = QCheckBox("%s" % header[j])
            self.formLayout.addRow(checkbox)

        groupBox.setLayout(self.formLayout)

        self.var_select.setWidget(groupBox)
        self.var_select.setWidgetResizable(True)


    def send_selected(self):
        self.conti_lst = []
        self.categorial_lst = []
        for i in range(self.formLayout.count()):
            chBox = self.formLayout.itemAt(i).widget()
            if chBox.isChecked():
                self.categorial_lst.append(chBox.text())
            else:
                self.conti_lst.append(chBox.text())

        self.col_value_cut.clear()
        self.col_value_cut.addItems(self.conti_lst)

        self.col_value_cut.setEnabled(True)
        self.cut_btn.setEnabled(True)
        self.cut_input.setEnabled(True)


        content = '進行變數類別選擇\n類別變數：'
        for ele in self.categorial_lst:
            if ele == self.categorial_lst[:-1]:
                content = content + ele
            else:
                content = content + ele + ', '
        content = content + '\n連續變數：'
        for ele in self.conti_lst:
            if ele == self.conti_lst[:-1]:
                content = content + ele
            else:
                content = content + ele + ', '
        self.document.add_paragraph(content, style='Intense Quote')


    def cut(self):
        self.cut_col = str(self.col_value_cut.currentText())
        cut_str = self.cut_input.text()
        try:
            cut_list = cut_str.split()
            cut_list = list(map(float, cut_list))

            self.grouped_data = pre_processing.group_data(self.df, self.cut_col, cut_list)

            self.renewal()

            self.col_value_cut.setEnabled(False)
            self.cut_input.setEnabled(False)
            self.cut_btn.setEnabled(False)
        except:
            QMessageBox.about(self, "輸入格式錯誤", "請依照系統指示格式進行切斷點輸入")
            self.cut_input.clear()

    def show_cate_plot(self):
        bars, height, file_path = des_sta.show_grouped_des(self.grouped_data, self.cut_col)

        self.formLayout = QFormLayout()
        groupBox = QGroupBox()

        for i in range(len(bars)):
            content = str("{}: {}".format(bars[i], height[i]))
            label = QLabel(content)
            self.formLayout.addRow(label)
        label_img = QPixmap(file_path).scaled(500,500, aspectRatioMode=True)
        label = QLabel()
        label.setPixmap(label_img)
        self.formLayout.addRow(label)

        groupBox.setLayout(self.formLayout)

        self.group_outcome.setWidget(groupBox)
        self.group_outcome.setWidgetResizable(True)

        self.send_cut_btn.setEnabled(True)
        self.cancel_btn.setEnabled(True)

    def cancel_cut(self):
        self.renew_group_outcome()

        self.col_value_cut.setEnabled(True)
        self.cut_input.setEnabled(True)
        self.cut_btn.setEnabled(True)

    def send_cut(self):
        self.renew_group_outcome()

        self.df[self.cut_col] = self.grouped_data
        self.conti_lst.remove(self.cut_col)
        self.categorial_lst.append(self.cut_col)

        self.col_value_cut.setEnabled(True)
        self.cut_input.setEnabled(True)
        self.cut_btn.setEnabled(True)

        content = '進行資料分組，分組變數為：' + self.cut_col + '\n下方圖片顯示分組結果：'
        self.document.add_paragraph(content, style='Intense Quote')
        filepath = "img/"+self.cut_col+"_des.png"
        self.document.add_picture(filepath)

        self.renewal()

    def renew_group_outcome(self):
        self.cut_input.clear()
        for i in reversed(range(self.formLayout.count())):
            self.formLayout.itemAt(i).widget().setParent(None)
        groupBox = QGroupBox()
        groupBox.setLayout(self.formLayout)

        self.group_outcome.setWidget(groupBox)
        self.group_outcome.setWidgetResizable(True)


    #tab4
    def show_des(self):
        if self.conti_lst != []:
            descriptive_sta = self.df[self.conti_lst].describe()
            self.df_des = descriptive_sta.iloc[[0, 1, 2, 3, 5, 7]]
            self.df_des = self.df_des.round(4)
            model_des = des_sta.pandasModel(self.df_des)
            self.desTable_conti.setModel(model_des)

            formLayout_conti = QFormLayout()
            groupBox_conti = QGroupBox()
            for i in self.conti_lst:
                label = QLabel(i)
                formLayout_conti.addRow(label)
                file_path = des_sta.show_conti_des(self.df, i)
                label_img = QPixmap(file_path).scaled(500, 500, aspectRatioMode=True)
                label = QLabel()
                label.setPixmap(label_img)
                formLayout_conti.addRow(label)

            groupBox_conti.setLayout(formLayout_conti)

            self.conti_des_area.setWidget(groupBox_conti)
            self.conti_des_area.setWidgetResizable(True)

        if self.categorial_lst != []:
            formLayout_cate = QFormLayout()
            groupBox_cate = QGroupBox()
            for i in self.categorial_lst:
                bars, height, file_path = des_sta.show_cate_des(self.df, i)
                label = QLabel(i)
                formLayout_cate.addRow(label)

                # to change
                sum = 0
                for i in height:
                    sum += i

                for i in range(len(bars)):
                    pa = height[i] / sum * 100
                    content = str("{}: {}({:,.4f}%)".format(bars[i], height[i], pa))
                    label = QLabel(content)
                    formLayout_cate.addRow(label)
                label_img = QPixmap(file_path).scaled(500, 500, aspectRatioMode=True)
                label = QLabel()
                label.setPixmap(label_img)
                formLayout_cate.addRow(label)

            groupBox_cate.setLayout(formLayout_cate)

            self.cate_des_area.setWidget(groupBox_cate)
            self.cate_des_area.setWidgetResizable(True)






    #tab5
    def convert(self):
        record = []
        for i in self.column_list:
            if self.df[i].dtype == "object":
                self.df[i] = pre_processing.str_to_num(self.df, i)
                self.renewal()
                record.append(i)
        if record != []:
            content = '執行演算法運算前，系統已自動將文字內容轉換為數值'
            for ele in record:
                if ele == record[:-1]:
                    content = content + ele
                else:
                    content = content + ele + ', '
            self.document.add_paragraph(content, style='Intense Quote')

    def cart_show(self):
        #add descriptive statistic into document
        self.document.add_heading('二、敘述統計結果', level=1)
        for i in self.column_list:
            content = i + ": "
            self.document.add_paragraph(content, style='Intense Quote')
            filepath = "img/"+i+"_des.png"
            try:
                self.document.add_picture(filepath)
            except:
                pass
        #add algoeithms results into document
        self.document.add_heading('三、演算法結果', level=1)
        if self.cart_CC.text() != "" and self.cart_RS.text!= "":
            cost_complexity = self.cart_CC.text()
            random_state = self.cart_RS.text()

            y_test, pre, acc, sens_score, spec, auc, precision, npv_score, f1score = cart.training(self.df, cost_complexity, random_state)#still to change
            filename1 = cart.conf_matrix(y_test, pre)
            filename2 = cart.roccurve(y_test, pre)

            confusion_f = QPixmap(filename1)
            roc_f = QPixmap(filename2)

            self.cart_acc.setText("%.4f"%acc)
            self.cart_sen.setText("%.4f"%sens_score)
            self.cart_spe.setText("%.4f"%spec)
            self.cart_con.setPixmap(confusion_f)
            self.cart_roc.setPixmap(roc_f)
            self.cart_auc.setText("%.4f"%auc)
            self.cart_precision.setText("%.4f" % precision)
            self.cart_f1.setText("%.4f" % f1score)
            self.cart_NPV.setText("%.4f" % npv_score)

            self.tabWidget.setCurrentWidget(self.tab05)

            p = self.document.add_paragraph("CART: ", style='Intense Quote')
            content = "正確率: " + str(acc) + "、敏感度: " + str(sens_score) + "、特異度: " + str(spec) + "\n"
            p.add_run(content)
            content = "precision: " + str(precision) + "、f1 score: " + str(f1score) + "、NPV: " + str(npv_score) + "、auc: " + str(auc)
            p.add_run(content)
            self.document.add_picture("img/cart_confu.png")
            self.document.add_picture("img/cart_roc.png")
        else:
            pass



    def knn_show(self):
        if self.k_value.text() != "":
            k_value = self.k_value.text()

            y_test, pre, acc, sens_score, spec, auc, precision, npv_score, f1score = knn.training(self.df, k_value)#still to change
            filename1 = knn.conf_matrix(y_test, pre)
            filename2 = knn.roccurve(y_test, pre)

            confusion_f = QPixmap(filename1)
            roc_f = QPixmap(filename2)

            self.knn_acc.setText("%.4f" % acc)
            self.knn_sen.setText("%.4f" % sens_score)
            self.knn_spe.setText("%.4f" % spec)
            self.knn_con.setPixmap(confusion_f)
            self.knn_roc.setPixmap(roc_f)
            self.knn_auc.setText("%.4f" % auc)
            self.knn_precision.setText("%.4f" % precision)
            self.knn_f1.setText("%.4f" % f1score)
            self.knn_NPV.setText("%.4f" % npv_score)

            p = self.document.add_paragraph("KNN: ", style='Intense Quote')
            content = "正確率: " + str(acc) + "、敏感度: " + str(sens_score) + "、特異度: " + str(spec) + "\n"
            p.add_run(content)
            content = "precision: " + str(precision) + "、f1 score: " + str(f1score) + "、NPV: " + str(npv_score) + "、auc: " + str(auc)
            p.add_run(content)
            self.document.add_picture("img/knn_confu.png")
            self.document.add_picture("img/knn_roc.png")
        else:
            pass



    def rf_show(self):
        if self.rf_MD.text() != "":
            max_depth = self.rf_MD.text()

            y_test, pre, acc, sens_score, spec, auc, precision, npv_score, f1score = rf.training(self.df, max_depth)
            filename1 = rf.conf_matrix(y_test, pre)
            filename2 = rf.roccurve(y_test, pre)

            confusion_f = QPixmap(filename1)
            roc_f = QPixmap(filename2)

            self.rf_acc.setText("%.4f" % acc)
            self.rf_sen.setText("%.4f" % sens_score)
            self.rf_spe.setText("%.4f" % spec)
            self.rf_con.setPixmap(confusion_f)
            self.rf_roc.setPixmap(roc_f)
            self.rf_auc.setText("%.4f"%auc)
            self.rf_precision.setText("%.4f" % precision)
            self.rf_f1.setText("%.4f" % f1score)
            self.rf_NPV.setText("%.4f" % npv_score)

            p = self.document.add_paragraph("Random Forest: ", style='Intense Quote')
            content = "正確率: " + str(acc) + "、敏感度: " + str(sens_score) + "、特異度: " + str(spec) + "\n"
            p.add_run(content)
            content = "precision: " + str(precision) + "、f1 score: " + str(f1score) + "、NPV: " + str(npv_score) + "、auc: " + str(auc)
            p.add_run(content)
            self.document.add_picture("img/rf_confu.png")
            self.document.add_picture("img/rf_roc.png")
        else:
            pass

    def svm_show(self):
        if self.svm_C.text() != "":
            c = self.svm_C.text()

            y_test, pre, acc, sens_score, spec, auc, precision, npv_score, f1score = svm.training(self.df, c)
            filename1 = svm.conf_matrix(y_test, pre)
            filename2 = svm.roccurve(y_test, pre)

            confusion_f = QPixmap(filename1)
            roc_f = QPixmap(filename2)

            self.svm_acc.setText("%.4f" % acc)
            self.svm_sen.setText("%.4f" % sens_score)
            self.svm_spe.setText("%.4f" % spec)
            self.svm_con.setPixmap(confusion_f)
            self.svm_roc.setPixmap(roc_f)
            self.svm_auc.setText("%.4f" % auc)
            self.svm_precision.setText("%.4f" % precision)
            self.svm_f1.setText("%.4f" % f1score)
            self.svm_NPV.setText("%.4f" % npv_score)

            p = self.document.add_paragraph("SVM: ", style='Intense Quote')
            content = "正確率: " + str(acc) + "、敏感度: " + str(sens_score) + "、特異度: " + str(spec) + "\n"
            p.add_run(content)
            content = "precision: " + str(precision) + "、f1 score: " + str(f1score) + "、NPV: " + str(npv_score) + "、auc: " + str(auc)
            p.add_run(content)
            self.document.add_picture("img/svm_confu.png")
            self.document.add_picture("img/svm_roc.png")
        else:
            pass

    def logistic_show(self):
        if self.logistic_C.text() != "":
            c = self.logistic_C.text()

            y_test, pre, acc, sens_score, spec, auc, precision, npv_score, f1score = logistic.training(self.df, c)
            filename1 = logistic.conf_matrix(y_test, pre)
            filename2 = logistic.roccurve(y_test, pre)

            confusion_f = QPixmap(filename1)
            roc_f = QPixmap(filename2)

            self.logistic_acc.setText("%.4f" % acc)
            self.logistic_sen.setText("%.4f" % sens_score)
            self.logistic_spe.setText("%.4f" % spec)
            self.logistic_con.setPixmap(confusion_f)
            self.logistic_roc.setPixmap(roc_f)
            self.logistic_auc.setText("%.4f" % auc)
            self.logistic_precision.setText("%.4f" % precision)
            self.logistic_f1.setText("%.4f" % f1score)
            self.logistic_NPV.setText("%.4f" % npv_score)

            p = self.document.add_paragraph("Logistic: ", style='Intense Quote')
            content = "正確率: " + str(acc) + "、敏感度: " + str(sens_score) + "、特異度: " + str(spec) + "\n"
            p.add_run(content)
            content = "precision: " + str(precision) + "、f1 score: " + str(f1score) + "、NPV: " + str(npv_score) + "、auc: " + str(auc)
            p.add_run(content)
            self.document.add_picture("img/logistic_confu.png")
            self.document.add_picture("img/logistic_roc.png")
        else:
            pass

    def parameter_clear(self):
        self.cart_CC.clear()
        self.cart_RS.clear()
        self.k_value.clear()
        self.rf_MD.clear()
        self.svm_C.clear()
        self.logistic_C.clear()



    #tab6
    def download_csv(self):
        dialog = QFileDialog()
        foo_dir, _ = dialog.getSaveFileName(self, filter="CSV(*.csv)")
        try:
            file_path = foo_dir
            self.df.to_csv(file_path, index=False)
        except:
            pass

    def download_file(self):
        dialog = QFileDialog()
        foo_dir, _ = dialog.getSaveFileName(self, filter="Document(*.docx)")
        self.document.add_page_break()
        try:
            self.document.save(foo_dir)
        except:
            pass



    #else
    def reset(self):
        if self.fileName1 != '':
            self.df = pd.read_csv(self.fileName1)
            self.document = Document()
            self.document.add_heading('結果報表', 0)
            self.document.add_heading('一、前處理步驟', level=1)
            self.renewal()

            df_empty = pd.DataFrame({'': []})
            model = tableView.pandasModel(df_empty)
            self.desTable_conti.setModel(model)

            formLayout_conti = QFormLayout()
            groupBox_conti = QGroupBox()
            groupBox_conti.setLayout(formLayout_conti)
            self.conti_des_area.setWidget(groupBox_conti)

            formLayout_cate = QFormLayout()
            groupBox_cate = QGroupBox()
            groupBox_cate.setLayout(formLayout_cate)
            self.cate_des_area.setWidget(groupBox_cate)

            self.algorithm_empty()

            self.var_selection()

    def renewal(self):
        self.column_list = list(self.df.columns)

        model = tableView.pandasModel(self.df)
        self.dataTable.setModel(model)

        self.column_list = list(self.df.columns)
        self.col_value_delete.clear()
        self.col_value_delete.addItems(self.column_list)
        self.col_value_cut.clear()
        self.col_value_cut.addItems(self.conti_lst)

        null_img = pre_processing.show_missing(self.df)
        image = QPixmap(null_img)
        self.null_img.setPixmap(image)

        self.show_des()

    def showDialog(self, text, title):
        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)
        msgBox.setText(text)
        msgBox.setWindowTitle(title)
        msgBox.setStandardButtons(QMessageBox.Ok)

        returnValue = msgBox.exec()

    def closeEvent(self, a0: QtGui.QCloseEvent) -> None:

        folder = 'img'
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

    def algorithm_empty(self):
        self.cart_acc.clear()
        self.cart_con.clear()
        self.cart_roc.clear()
        self.cart_sen.clear()
        self.cart_spe.clear()
        self.cart_auc.clear()
        self.cart_precision.clear()
        self.cart_f1.clear()
        self.cart_NPV.clear()
        self.knn_acc.clear()
        self.knn_con.clear()
        self.knn_roc.clear()
        self.knn_sen.clear()
        self.knn_spe.clear()
        self.knn_auc.clear()
        self.knn_precision.clear()
        self.knn_f1.clear()
        self.knn_NPV.clear()
        self.rf_acc.clear()
        self.rf_con.clear()
        self.rf_roc.clear()
        self.rf_sen.clear()
        self.rf_spe.clear()
        self.rf_auc.clear()
        self.rf_precision.clear()
        self.rf_f1.clear()
        self.rf_NPV.clear()
        self.svm_acc.clear()
        self.svm_con.clear()
        self.svm_roc.clear()
        self.svm_sen.clear()
        self.svm_spe.clear()
        self.svm_auc.clear()
        self.svm_precision.clear()
        self.svm_f1.clear()
        self.svm_NPV.clear()
        self.logistic_acc.clear()
        self.logistic_con.clear()
        self.logistic_roc.clear()
        self.logistic_sen.clear()
        self.logistic_spe.clear()
        self.logistic_auc.clear()
        self.logistic_precision.clear()
        self.logistic_f1.clear()
        self.logistic_NPV.clear()







if __name__ == '__main__':
    app = QtWidgets.QApplication([])
    app.setStyle('Fusion')
    gui = UI()
    sys.exit(app.exec_())




