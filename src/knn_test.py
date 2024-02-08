from PyQt5.QtWidgets import *
import sys,pickle,os
from PyQt5 import uic, QtWidgets ,QtCore, QtGui
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsClassifier as KNC
from sklearn import metrics
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from sklearn.metrics import accuracy_score, f1_score, roc_curve, auc
import os
import docx

class UI(QMainWindow):
    def __init__(self):
        super(UI, self).__init__()
        uic.loadUi("knn_test.ui", self)

        # self.count = 0
        self.browse = self.findChild(QPushButton,"browse")
        self.acc = self.findChild(QLabel,"acc")
        self.sens = self.findChild(QLabel,"sensitivity")
        self.f1score = self.findChild(QLabel,"f1")
        self.spec = self.findChild(QLabel,"specificity")
        self.neighbor = self.findChild(QLineEdit, "neighbor")
        self.conf_bt = self.findChild(QPushButton,"conf_mat")
        self._roc = self.findChild(QPushButton,"roc_curve")
        self.train = self.findChild(QPushButton,"train")
        self.browse.clicked.connect(self.getCSV)
        self.train.clicked.connect(self.training)
        self.conf_bt.clicked.connect(self.conf_matrix)
        self._roc.clicked.connect(self.roccurve)
        self.show()

    def getCSV(self):
        self.filePath, _ = QtWidgets.QFileDialog.getOpenFileName(self, 'Open file', '/Users/hankchang/Desktop',"csv(*.csv)")
        if(self.filePath):
            self.df = pd.read_csv(str(self.filePath))
            self.df.dropna(axis=0, inplace=True)
            self.df.reset_index(drop=True, inplace=True)



    def training(self):
        self.x_train, self.x_test, self.y_train, self.y_test = train_test_split(self.df.iloc[:,0:-1],self.df.iloc[:,-1],\
        train_size=0.8,test_size=0.2,random_state=24)
        self.lr = KNC(n_neighbors=int(self.neighbor.text()))
        self.lr.fit(self.x_train, self.y_train)
        self.pre = self.lr.predict(self.x_test)
        cm = metrics.confusion_matrix(self.y_test, self.pre)
        sens_score = (cm[0,0]/(cm[0,0] + cm[0,1])).round(6)
        spec_score = (cm [1,1]/(cm[1,0] + cm[1,1])).round(6)
        # precision(PPV) and NPV
        precision = (cm[0,0]/(cm[0,0] + cm[1,0])).round(6)
        npv_score = (cm[1,1]/(cm[0,1] + cm[1,1])).round(6)
        #auc value
        fpr, tpr, _ = roc_curve(self.y_test, self.pre)
        roc_auc = auc(fpr, tpr)

        self.acc.setText(str(accuracy_score(self.y_test, self.pre).round(6)))
        self.sens.setText(str(sens_score))
        self.f1score.setText(str(f1_score(self.y_test, self.pre).round(6)))
        self.spec.setText(str(spec_score))
        self.precision.setText(str(precision))
        self.npv.setText(str(npv_score))
        #setText auc_QLabel
        self.auc_label.setText(str(roc_auc))
    def conf_matrix(self):
        # self.count += 1
        data = {'y_Actual':self.y_test, 'y_Predicted':self.pre }
        df = pd.DataFrame(data, columns=['y_Actual','y_Predicted'])
        confusion_matrix = pd.crosstab(df['y_Actual'], df['y_Predicted'], rownames=['Actual'], colnames=['Predicted'])
        self.con_fig = plt.figure()
        plt.title("Confusion Matrix")
        sns.heatmap(confusion_matrix, annot=True,linewidths=.5)
        plt.savefig('con_matrix.png')
        # pip install python-docx , import docx
                # if(os.path.exists('report.docx')):
                #     doc = docx.Document('report.docx')
                #     doc.add_paragraph(f'This is your {self.count} times open the file')
                #     doc.save('report.docx')  # save docx
                # else:
        doc = docx.Document()
        doc.add_heading('Confusion Matrix', level=1)  # adding head
        doc.add_picture("con_matrix.png")  # add picture
        doc.save('report.docx')  # save docx
        os.remove('con_matrix.png') # delete png
        self.con_fig.show()


    def roccurve(self):
        fpr, tpr, threshold = roc_curve(self.y_test, self.pre)
        roc_auc = auc(fpr, tpr)
        plt.title('Receiver Operating Characteristic')
        plt.plot(fpr, tpr, 'b', label='AUC = %0.2f' % roc_auc)
        plt.legend(loc='lower right')
        plt.plot([0, 1], [0, 1], 'r--')
        plt.xlim([0, 1])
        plt.ylim([0, 1])
        plt.ylabel('True Positive Rate')
        plt.xlabel('False Positive Rate')
        plt.show()


app = QApplication(sys.argv)
window = UI()
app.exec_()